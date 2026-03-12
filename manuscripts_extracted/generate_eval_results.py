from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading.append(elm)

def setup_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(1.5)
        s.right_margin = Cm(1.5)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    return doc

def add_p(doc, text, bold=False, size=12, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

# Evaluation results from Run 4
results = [
    {"id": 1,  "input": "Bachelor of Science in Information Technology",              "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Database Management", "System Administration", "Networking", "IT Support", "Software Development"], "fp": [], "fn": []},
    {"id": 2,  "input": "Bachelor of Science in Accountancy",                         "p": 0.83, "r": 1.00, "f1": 0.91, "status": "PASS",
     "tp": ["Accounting", "Auditing", "Taxation", "Financial Reporting", "Financial Analysis"], "fp": ["Bookkeeping"], "fn": []},
    {"id": 3,  "input": "Bachelor of Science in Computer Science",                    "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Programming", "Data Structures", "Algorithms", "Software Development", "Computer Architecture"], "fp": [], "fn": []},
    {"id": 4,  "input": "BS Business Administration major in Financial Management",   "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Financial Analysis", "Financial Management", "Investment Management", "Corporate Finance", "Business Administration"], "fp": [], "fn": []},
    {"id": 5,  "input": "Bachelor of Science in Criminology",                         "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Criminal Justice", "Forensic Science", "Criminology", "Legal Procedures", "Investigative Techniques"], "fp": [], "fn": []},
    {"id": 6,  "input": "Full-Stack Web Development Bootcamp",                        "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Web Development", "Frontend Development", "Backend Development", "Database Management", "API Development", "JavaScript"], "fp": [], "fn": []},
    {"id": 7,  "input": "AWS Cloud Practitioner Bootcamp",                             "p": 0.83, "r": 1.00, "f1": 0.91, "status": "PASS",
     "tp": ["AWS", "Cloud Computing", "Cloud Security", "Cloud Architecture", "Scalability"], "fp": ["Cost Management"], "fn": []},
    {"id": 8,  "input": "UI/UX Design Intensive Bootcamp",                             "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["UI/UX Design", "Prototyping", "User Research", "Wireframing", "Usability Testing"], "fp": [], "fn": []},
    {"id": 9,  "input": "Cybersecurity Fundamentals Bootcamp",                         "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Cybersecurity", "Network Security", "Information Security", "Threat Analysis", "Incident Response"], "fp": [], "fn": []},
    {"id": 10, "input": "Python",                                                      "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Python", "Programming", "Scripting"], "fp": [], "fn": []},
    {"id": 11, "input": "Regional Hackathon 2026",                                     "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Problem Solving", "Programming", "Teamwork", "Prototyping", "Software Development"], "fp": [], "fn": []},
    {"id": 12, "input": "National Cybersecurity CTF Competition 2025",                 "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Cybersecurity", "Network Security", "Cryptography", "Problem Solving", "Penetration Testing"], "fp": [], "fn": []},
    {"id": 13, "input": "University Data Science Summit — Best Presenter Award",       "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Data Science", "Public Speaking", "Presentation Skills", "Data Visualization", "Research"], "fp": [], "fn": []},
    {"id": 14, "input": "Fintech Innovation Challenge 2025 — Finalist",                "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Financial Technology", "Innovation", "Business Strategy", "Problem Solving", "Prototyping"], "fp": [], "fn": []},
    {"id": 15, "input": "TESDA NC II in Computer Systems Servicing",                   "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Computer Hardware Repair", "Network Configuration", "Operating System Installation", "Troubleshooting", "Preventive Maintenance"], "fp": [], "fn": []},
    {"id": 16, "input": "TESDA NC III in Bookkeeping",                                 "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Bookkeeping", "Accounting", "Financial Record Keeping", "Financial Reporting", "Payroll"], "fp": [], "fn": []},
    {"id": 17, "input": "TESDA NC II in Bread and Pastry Production",                  "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Baking", "Food Safety", "Pastry Making", "Bread Making", "Kitchen Operations"], "fp": [], "fn": []},
    {"id": 18, "input": "TESDA NC II in Electrical Installation and Maintenance",      "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Electrical Installation", "Electrical Maintenance", "Electrical Wiring", "Troubleshooting", "Electrical Safety"], "fp": [], "fn": []},
    {"id": 19, "input": "DICT Digital Literacy Certificate",                           "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Digital Literacy", "Computer Fundamentals", "Online Safety", "Internet Navigation", "Productivity Software"], "fp": [], "fn": []},
    {"id": 20, "input": "DICT Cybersecurity Essentials Certification",                 "p": 1.00, "r": 1.00, "f1": 1.00, "status": "PASS",
     "tp": ["Cybersecurity", "Network Security", "Information Security", "Risk Management", "Incident Response"], "fp": [], "fn": []},
]

doc = setup_doc()

# Title
h = doc.add_heading('VECTOR — AI Skill Extractor F1 Evaluation Results', level=1)
for r in h.runs:
    r.font.name = 'Times New Roman'

add_p(doc, 'Evaluation Run 4 — Final Results', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, '20 test cases · Match mode: SOFT · Model: Gemini 2.5 Flash', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(100,100,100))

doc.add_paragraph()

# Summary Box
add_p(doc, 'Aggregate Results', bold=True, size=13)
summary_table = doc.add_table(rows=2, cols=6)
summary_table.style = 'Table Grid'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

summary_headers = ['Precision', 'Recall', 'F1 Score', 'PASS', 'PARTIAL', 'FAIL']
summary_values = ['0.99', '1.00', '0.99', '20', '0', '0']

for i, h_text in enumerate(summary_headers):
    cell = summary_table.rows[0].cells[i]
    cell.text = h_text
    set_cell_shading(cell, '1c4587')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
            run.bold = True

for i, val in enumerate(summary_values):
    cell = summary_table.rows[1].cells[i]
    cell.text = val
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            if i < 3:
                run.font.color.rgb = RGBColor(0, 128, 0)
            elif i == 3:
                run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()

# Per-case results table
add_p(doc, 'Per-Case Results', bold=True, size=13)

case_table = doc.add_table(rows=21, cols=6)
case_table.style = 'Table Grid'
case_table.alignment = WD_TABLE_ALIGNMENT.CENTER

case_headers = ['#', 'Credential Input', 'Precision', 'Recall', 'F1', 'Status']
for i, h_text in enumerate(case_headers):
    cell = case_table.rows[0].cells[i]
    cell.text = h_text
    set_cell_shading(cell, '1c4587')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255,255,255)
            run.bold = True

for ri, r in enumerate(results):
    row = case_table.rows[ri + 1]
    row.cells[0].text = str(r['id'])
    row.cells[1].text = r['input'][:55]
    row.cells[2].text = f"{r['p']:.2f}"
    row.cells[3].text = f"{r['r']:.2f}"
    row.cells[4].text = f"{r['f1']:.2f}"
    row.cells[5].text = r['status']

    for ci, cell in enumerate(row.cells):
        if ri % 2 == 1:
            set_cell_shading(cell, 'f0f4f8')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(8)
                if ci == 5:
                    run.font.color.rgb = RGBColor(0, 128, 0)
                    run.bold = True

doc.add_page_break()

# Detailed extraction results
add_p(doc, 'Detailed Extraction Results per Case', bold=True, size=13)

# Group by type
groups = {
    'Academic Degrees': results[0:5],
    'Bootcamp Certificates': results[5:10],
    'Event Badges': results[10:14],
    'Government Certifications': results[14:20],
}

for group_name, group_results in groups.items():
    add_p(doc, group_name, bold=True, size=11)
    
    for r in group_results:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        
        run = p.add_run(f"Case {r['id']}: ")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = True
        
        run = p.add_run(r['input'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        
        run = p.add_run(f"  — F1: {r['f1']:.2f} ")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 128, 0)
        run.bold = True
        
        run = p.add_run(f"[{r['status']}]")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 128, 0)
        run.bold = True
        
        # TP
        if r['tp']:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(1)
            p2.paragraph_format.left_indent = Cm(1)
            run = p2.add_run("✓ Matched: ")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 128, 0)
            run.bold = True
            run = p2.add_run(", ".join(r['tp']))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
        
        # FP
        if r['fp']:
            p3 = doc.add_paragraph()
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after = Pt(1)
            p3.paragraph_format.left_indent = Cm(1)
            run = p3.add_run("✗ Extra (FP): ")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(200, 0, 0)
            run.bold = True
            run = p3.add_run(", ".join(r['fp']))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
        
        # FN
        if r['fn']:
            p4 = doc.add_paragraph()
            p4.paragraph_format.space_before = Pt(0)
            p4.paragraph_format.space_after = Pt(1)
            p4.paragraph_format.left_indent = Cm(1)
            run = p4.add_run("✗ Missed (FN): ")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(200, 150, 0)
            run.bold = True
            run = p4.add_run(", ".join(r['fn']))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)

doc.add_paragraph()
add_p(doc, 'Methodology', bold=True, size=13)
add_p(doc, 'The evaluation was conducted against a golden dataset of 20 test cases spanning 4 credential types: academic degrees (5), bootcamp certificates (5), event badges (4), and government certifications (6). Each case was processed by the Gemini 2.5 Flash NLP model using the VECTOR skill extraction prompt.', size=10)
add_p(doc, 'Matching Strategy: SOFT match — extracted tags are accepted if they contain or are contained by expected tags after normalization (lowercase, trim, replace hyphens/underscores with spaces). This handles variations like "machine-learning" ↔ "Machine Learning".', size=10)
add_p(doc, 'Metrics: Micro-averaged Precision, Recall, and F1-score computed across all 20 cases.', size=10)

doc.save('C:/vector-system/VECTOR_F1_Evaluation_Results.docx')
print('Saved: VECTOR_F1_Evaluation_Results.docx')

# Also generate clean text version
lines = []
lines.append('=' * 70)
lines.append('VECTOR — AI Skill Extractor F1 Evaluation Results')
lines.append('Evaluation Run 4 — Final Results')
lines.append('20 cases · Match mode: SOFT · Model: Gemini 2.5 Flash')
lines.append('=' * 70)
lines.append('')
lines.append(f"{'AGGREGATE RESULTS':^70}")
lines.append(f"  Precision: 0.99    Recall: 1.00    F1: 0.99")
lines.append(f"  20/20 PASS | 0 PARTIAL | 0 FAIL | 0 ERROR")
lines.append('')
lines.append('-' * 70)
lines.append(f"{'#':<4} {'Credential Input':<50} {'P':>5} {'R':>5} {'F1':>5} {'Status':>7}")
lines.append('-' * 70)

for r in results:
    lines.append(f"{r['id']:<4} {r['input'][:50]:<50} {r['p']:>5.2f} {r['r']:>5.2f} {r['f1']:>5.2f} {r['status']:>7}")

lines.append('-' * 70)
lines.append('')
lines.append('DETAILED EXTRACTION RESULTS')
lines.append('=' * 70)

for group_name, group_results in groups.items():
    lines.append('')
    lines.append(f'--- {group_name} ---')
    for r in group_results:
        lines.append(f"  Case {r['id']}: {r['input']}")
        lines.append(f"    F1: {r['f1']:.2f} [{r['status']}]")
        if r['tp']:
            lines.append(f"    ✓ Matched: {', '.join(r['tp'])}")
        if r['fp']:
            lines.append(f"    ✗ Extra:   {', '.join(r['fp'])}")
        if r['fn']:
            lines.append(f"    ✗ Missed:  {', '.join(r['fn'])}")

lines.append('')
lines.append('=' * 70)

with open('C:/vector-system/eval_results_4.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))
print('Saved: eval_results_4.txt (clean)')
